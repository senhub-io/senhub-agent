#!/usr/bin/env python3
"""
Apply professional visual styling to Word document:
- Color headings (blue #28b3d0 for H1, gradient for others)
- Professional fonts (Calibri for body, Arial for headings)
- Improved spacing and layout
- Page breaks for H1
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_heading_style(paragraph, level, add_page_break=True):
    """Apply professional styling to headings with automatic numbering"""

    # Color scheme (SenHub blue)
    colors = {
        1: RGBColor(40, 179, 208),   # #28b3d0 - SenHub blue
        2: RGBColor(51, 51, 51),      # Dark gray
        3: RGBColor(102, 102, 102),   # Medium gray
        4: RGBColor(128, 128, 128),   # Light gray
        5: RGBColor(153, 153, 153),   # Lighter gray
        6: RGBColor(153, 153, 153),   # Lighter gray
    }

    # Font sizes
    sizes = {
        1: 20,
        2: 16,
        3: 14,
        4: 12,
        5: 11,
        6: 11,
    }

    # Left alignment for headings
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set font - Roboto for headings
    for run in paragraph.runs:
        run.font.name = 'Roboto'
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
        run.font.bold = True

    # Spacing
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)

    spacing = OxmlElement('w:spacing')
    if level == 1:
        spacing.set(qn('w:before'), '720')  # 36pt before H1 (more space)
        spacing.set(qn('w:after'), '360')   # 18pt after H1
    elif level == 2:
        spacing.set(qn('w:before'), '480')  # 24pt before H2
        spacing.set(qn('w:after'), '240')   # 12pt after H2
    else:
        spacing.set(qn('w:before'), '360')  # 18pt before H3+
        spacing.set(qn('w:after'), '180')   # 9pt after H3+
    pPr.append(spacing)

    # Add page break before H1 using pageBreakBefore property
    if level == 1 and add_page_break:
        # Use pageBreakBefore property in XML (more reliable than manual breaks)
        pageBreak = pPr.find(qn('w:pageBreakBefore'))
        if pageBreak is None:
            pageBreak = OxmlElement('w:pageBreakBefore')
            # Insert at the beginning of pPr for proper positioning
            pPr.insert(0, pageBreak)

def set_body_style(paragraph):
    """Apply professional styling to body text"""
    for run in paragraph.runs:
        if run.font.name != 'Consolas':  # Don't change code font
            run.font.name = 'Roboto'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Left alignment
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Line spacing
    paragraph.paragraph_format.line_spacing = 1.15

    # Paragraph spacing
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240')  # 12pt after paragraphs (more space)
    pPr.append(spacing)

def set_list_style(paragraph, is_first_item=False):
    """Apply spacing to list items"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)

    spacing = OxmlElement('w:spacing')

    # Add extra space before first list item (separates from preceding paragraph)
    if is_first_item:
        spacing.set(qn('w:before'), '180')  # 9pt before first item

    spacing.set(qn('w:after'), '120')  # 6pt after each list item (balanced spacing)
    pPr.append(spacing)

    # Set font
    for run in paragraph.runs:
        if run.font.name != 'Consolas':
            run.font.name = 'Roboto'
            run.font.size = Pt(11)


def style_title_page(doc):
    """Create professional SenHub title page"""
    if len(doc.paragraphs) >= 1:
        # Main title - Large, centered, SenHub blue
        title = doc.paragraphs[0]
        for run in title.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(32)
            run.font.color.rgb = RGBColor(40, 179, 208)  # SenHub blue
            run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing - push down from top
        pPr = title._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '1440')  # 72pt before (1 inch from top)
        spacing.set(qn('w:after'), '360')    # 18pt after
        pPr.append(spacing)

    if len(doc.paragraphs) >= 2:
        # Subtitle - Medium, centered, gray
        subtitle = doc.paragraphs[1]
        for run in subtitle.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
            run.font.bold = False
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        pPr = subtitle._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '720')  # 36pt after
        pPr.append(spacing)

    # Add SenHub branding at bottom (if not already exists)
    # Note: This would be in the footer ideally, but we add as last paragraph
    if len(doc.paragraphs) >= 3:
        # Author/company
        author = doc.paragraphs[2]
        for run in author.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(102, 102, 102)
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(doc.paragraphs) >= 4:
        # Date
        date_para = doc.paragraphs[3]
        for run in date_para.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(153, 153, 153)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def improve_table_appearance(table):
    """Improve table visual appearance"""
    # Style header row (first row)
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            # Gray background for header
            shading_elm = cell._element.get_or_add_tcPr().find(qn('w:shd'))
            if shading_elm is None:
                shading_elm = OxmlElement('w:shd')
                cell._element.get_or_add_tcPr().append(shading_elm)
            shading_elm.set(qn('w:fill'), 'E8E8E8')  # Light gray

            # Bold text in header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(51, 51, 51)

def main():
    if len(sys.argv) != 2:
        print("Usage: apply-visual-style.py <document.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    print(f"Applying visual styling: {docx_path}")

    try:
        doc = Document(docx_path)

        # 1. Style title page
        print("  → Styling title page...")
        style_title_page(doc)

        # 2. Style all paragraphs (this adds page breaks to H1)
        print("  → Styling paragraphs and headings...")
        heading_count = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
        first_h1_seen = False
        prev_was_list = False

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            # Skip code blocks (already styled)
            if style_name in ['Source Code', 'SourceCode', 'Verbatim', 'Code Block']:
                prev_was_list = False
                continue

            # Headings
            if 'Heading' in style_name:
                level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1

                # For H1: don't add page break to first one (comes after TOC)
                if level == 1:
                    add_break = first_h1_seen  # False for first H1, True for others
                    set_heading_style(paragraph, level, add_page_break=add_break)
                    if not first_h1_seen:
                        first_h1_seen = True
                        print(f"    → First H1 found (no page break)")
                    else:
                        print(f"    → H1 found (with page break)")
                else:
                    set_heading_style(paragraph, level, add_page_break=False)

                heading_count[level] = heading_count.get(level, 0) + 1
                prev_was_list = False
            # List items (including Pandoc's "Compact" style)
            elif 'List' in style_name or 'Compact' in style_name or paragraph._element.xpath('.//w:numPr'):
                is_first = not prev_was_list
                set_list_style(paragraph, is_first_item=is_first)
                prev_was_list = True
            else:
                # Body text
                set_body_style(paragraph)
                prev_was_list = False

        print(f"    ✓ Styled {sum(heading_count.values())} headings")

        # 3. Improve table appearance
        print(f"  → Improving {len(doc.tables)} table appearances...")
        for table in doc.tables:
            try:
                improve_table_appearance(table)
            except Exception as e:
                print(f"    ⚠ Warning: Could not style table: {e}")

        # Save
        doc.save(docx_path)
        print(f"  ✓ Visual styling applied successfully!")

    except Exception as e:
        print(f"  ✗ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
