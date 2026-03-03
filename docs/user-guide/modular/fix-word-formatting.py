#!/usr/bin/env python3
"""
Fix Word document formatting after Pandoc generation:
- Add visible borders to ALL tables
- Style code blocks with gray background
- Style inline code with light gray background
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_table_borders(table):
    """Add borders to table using XML manipulation"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()

    # Remove existing borders if any
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Create new borders element
    tblBorders = OxmlElement('w:tblBorders')

    # Define border for each side
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width (1/8 pt) - 4 = 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black border
        tblBorders.append(border)

    tblPr.append(tblBorders)

def optimize_column_widths(table):
    """Optimize column widths to minimize line wrapping"""
    if len(table.rows) == 0:
        return

    num_cols = len(table.rows[0].cells)

    # Calculate content length for each column
    col_max_lengths = [0] * num_cols

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < num_cols:
                # Calculate max text length in this cell
                cell_text = cell.text
                # Consider line breaks
                lines = cell_text.split('\n')
                max_line_length = max([len(line) for line in lines]) if lines else 0
                col_max_lengths[idx] = max(col_max_lengths[idx], max_line_length)

    # Calculate proportional widths
    total_length = sum(col_max_lengths)
    if total_length == 0:
        return

    # Available width (in inches) - leave margins
    available_width = Inches(6.5)  # ~6.5 inches for content

    # Set column widths proportionally
    for idx, max_length in enumerate(col_max_lengths):
        proportion = max_length / total_length
        width = available_width * proportion

        # Minimum width of 0.8 inches, maximum of 4 inches
        width = max(Inches(0.8), min(width, Inches(4.0)))

        # Apply to all cells in this column
        for row in table.rows:
            if idx < len(row.cells):
                row.cells[idx].width = width

def set_table_autofit(table):
    """Enable auto-fit for table to adjust to content"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()

    # Add table width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')
    tblPr.append(tblW)

    # Add table layout (auto)
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is not None:
        tblPr.remove(tblLayout)

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'auto')
    tblPr.append(tblLayout)

def set_cell_background(cell, color='FFFFFF'):
    """Set cell background color"""
    cell_xml = cell._element
    cell_properties = cell_xml.get_or_add_tcPr()

    # Remove existing shading
    shading = cell_properties.find(qn('w:shd'))
    if shading is not None:
        cell_properties.remove(shading)

    # Add new shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell_properties.append(shading)

def style_code_block(paragraph):
    """Style a code block paragraph"""
    # Set background color (light gray)
    pPr = paragraph._element.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is not None:
        pPr.remove(shd)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')  # Light gray
    pPr.append(shd)

    # Set paragraph spacing - increased for better separation
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')  # 12pt before (was 6pt)
    spacing.set(qn('w:after'), '240')   # 12pt after (was 6pt)
    pPr.append(spacing)

    # Style all runs in the paragraph
    for run in paragraph.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray text

def style_inline_code(run):
    """Style inline code run"""
    # Set background
    rPr = run._element.get_or_add_rPr()
    shd = rPr.find(qn('w:shd'))
    if shd is not None:
        rPr.remove(shd)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')  # Very light gray
    rPr.append(shd)

    # Set font
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 51, 51)

def is_code_paragraph(paragraph):
    """Check if paragraph is a code block"""
    style_name = paragraph.style.name if paragraph.style else ""

    # Pandoc code block styles
    code_styles = [
        'Source Code', 'SourceCode', 'Verbatim',
        'Code Block', 'Code', 'Literal'
    ]

    return style_name in code_styles

def is_code_run(run):
    """Check if run is inline code"""
    if not run.style:
        return False

    style_name = run.style.name
    code_styles = ['Verbatim Char', 'Source Text', 'Code', 'Inline Code']

    return style_name in code_styles

def align_table_cells_left(table):
    """Align all table cell contents to the left"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def main():
    if len(sys.argv) != 2:
        print("Usage: fix-word-formatting.py <document.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    print(f"Fixing formatting: {docx_path}")

    try:
        doc = Document(docx_path)

        # 1. Fix all tables
        print(f"  → Processing {len(doc.tables)} tables...")
        for i, table in enumerate(doc.tables):
            try:
                set_table_borders(table)
                optimize_column_widths(table)
                set_table_autofit(table)
                align_table_cells_left(table)
                print(f"    ✓ Table {i+1}: borders + optimized widths + left aligned")
            except Exception as e:
                print(f"    ✗ Table {i+1}: {e}")

        # 2. Fix code blocks in main document
        print("  → Styling code blocks...")
        code_blocks = 0
        for paragraph in doc.paragraphs:
            if is_code_paragraph(paragraph):
                style_code_block(paragraph)
                code_blocks += 1
        print(f"    ✓ Styled {code_blocks} code blocks")

        # 3. Fix inline code in main document
        print("  → Styling inline code...")
        inline_code = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if is_code_run(run):
                    style_inline_code(run)
                    inline_code += 1
        print(f"    ✓ Styled {inline_code} inline code snippets")

        # 4. Fix code in tables
        print("  → Styling code in tables...")
        table_code = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Code blocks in tables
                        if is_code_paragraph(paragraph):
                            style_code_block(paragraph)
                            table_code += 1
                        # Inline code in tables
                        for run in paragraph.runs:
                            if is_code_run(run):
                                style_inline_code(run)
                                table_code += 1

        if table_code > 0:
            print(f"    ✓ Styled {table_code} code elements in tables")

        # Save
        doc.save(docx_path)
        print(f"  ✓ Document saved successfully!")

    except Exception as e:
        print(f"  ✗ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
